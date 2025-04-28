from PyQt6 import QtCore, QtGui, QtWidgets
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import subprocess
import sys
import requests
from pathlib import Path
import tempfile
import psutil 
from docx.shared import Pt

APP_VERSION = "1.1.0"
VERSION_URL = "https://raw.githubusercontent.com/trungtien2410/hang-ra-quay/main/version.txt"
UPDATE_URL = "https://github.com/trungtien2410/hang-ra-quay/releases/download/V1.1.0/hang-ra-quay.exe"

def check_for_updates():
    try:
        response = requests.get(VERSION_URL, timeout=5)
        latest_version = response.text.strip()
        if latest_version != APP_VERSION:
            return latest_version
    except Exception as e:
        print(f"Update check failed: {e}")
    return None

def prompt_update(latest_version):
    reply = QtWidgets.QMessageBox.question(None, "Cập nhật mới",
        f"Có phiên bản mới {latest_version}. Bạn có muốn cập nhật không?",
        QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
    if reply == QtWidgets.QMessageBox.StandardButton.Yes:
        download_and_install()

def download_and_install():
    try:
        r = requests.get(UPDATE_URL, stream=True)
        if r.status_code == 200:
            total_size = int(r.headers.get('content-length', 0))
            block_size = 1024 * 1024
            temp_dir = tempfile.gettempdir()
            file_path = Path(temp_dir) / "update_installer.exe"

            progress = QtWidgets.QProgressDialog("🚀 Đang tải bản cập nhật...", "❌ Hủy", 0, 100)
            progress.setWindowModality(QtCore.Qt.WindowModality.ApplicationModal)
            progress.setWindowTitle("Cập nhật ứng dụng")
            progress.setFixedSize(400, 120)
            progress.setStyleSheet("""
                QProgressBar {
                    height: 24px;
                    border-radius: 12px;
                    background-color: #eeeeee;
                    border: 2px solid #cccccc;
                }
                QProgressBar::chunk {
                    background: qlineargradient(
                        x1: 0, y1: 0, x2: 1, y2: 0,
                        stop: 0 #4facfe, stop: 1 #00f2fe
                    );
                    border-radius: 12px;
                }
                QProgressDialog {
                    font-size: 14px;
                    font-family: "Segoe UI";
                }
            """)

            progress.setValue(0)

            with open(file_path, 'wb') as f:
                downloaded = 0
                start_time = time.time()
                for chunk in r.iter_content(block_size):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        percent = int(downloaded * 100 / total_size)
                        elapsed = time.time() - start_time
                        speed = downloaded / elapsed if elapsed > 0 else 0
                        remaining = (total_size - downloaded) / speed if speed > 0 else 0

                        progress.setLabelText(
                            f"🚀 Đang tải... {percent}%\n"
                            f"⬇ {downloaded // (1024 * 1024)} MB / {total_size // (1024 * 1024)} MB\n"
                            f"⏳ Còn lại: {int(remaining)}s"
                        )
                        progress.setValue(percent)
                        QtWidgets.QApplication.processEvents()
                        if progress.wasCanceled():
                            QtWidgets.QMessageBox.warning(None, "Hủy", "Đã hủy cập nhật.")
                            return

            QtWidgets.QMessageBox.information(None, "Hoàn tất", f"Đã tải xong!\n")
            QtWidgets.QApplication.quit()

    except Exception as e:
        QtWidgets.QMessageBox.critical(None, "Lỗi cập nhật", f"Không thể tải bản cập nhật: {e}")

# NEW FUNCTION to force close Word if file is open
def close_word_if_file_open(file_path):
    file_name = Path(file_path).name.lower()
    for proc in psutil.process_iter(['pid', 'name', 'open_files']):
        try:
            if proc.info['name'] and 'winword' in proc.info['name'].lower():
                open_files = proc.info['open_files']
                if open_files:
                    for opened in open_files:
                        if file_name in opened.path.lower():
                            proc.terminate()
                            proc.wait(timeout=3)
                            return True
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            pass
    return False

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return str(Path(sys._MEIPASS) / relative_path)
    else:
        return str(Path(__file__).parent / relative_path)

class Worker(QtCore.QThread):
    progress = QtCore.pyqtSignal(int)
    log = QtCore.pyqtSignal(str)
    finished = QtCore.pyqtSignal(object)

    def __init__(self, mnv, output_path):
        super().__init__()
        self.mnv = mnv
        self.output_path = output_path

    def run(self):
        self.log.emit("\U0001f300 Đang khởi tạo trình duyệt...")
        self.progress.emit(10)
        options = webdriver.ChromeOptions()
        options.add_argument('--headless=new')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-extensions')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--blink-settings=imagesEnabled=false') 

        
        browser = webdriver.Chrome(options=options)
        template_path = resource_path("test.docx")

        try:
            if not self.mnv.isnumeric():
                self.log.emit("❌ Mã nhân viên không hợp lệ.")
                self.progress.emit(0)
                self.finished.emit(None)
                return

            self.log.emit("🌐 Đang truy cập website...")
            browser.get("http://scfp.vn/Productscan.aspx")
            browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_txtMANV").send_keys(self.mnv)
            browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_rdbLoai_2").click()
            browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_txtXem").click()

            self.progress.emit(40)
            self.log.emit("📦 Đang lấy dữ liệu...")
            # time.sleep(1)
            WebDriverWait(browser, 5).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_ContentPlaceHolder1_gvPd"]/tbody'))
        )
            tbody = browser.find_element(By.XPATH, '//*[@id="ctl00_ContentPlaceHolder1_gvPd"]/tbody')
            rows = tbody.find_elements(By.TAG_NAME, "tr")
            data = [[cell.text for cell in row.find_elements(By.TAG_NAME, 'td')] for row in rows]

            df = pd.DataFrame(data[:], columns=data[0])
            df.iloc[:, 0] = pd.to_datetime(df.iloc[:, 0], format="%m/%d/%Y %I:%M:%S %p").dt.strftime('%d/%m/%Y')
            df.iloc[:, 3] = df.iloc[:, 4] + ' / ' + df.iloc[:, 3]
            df.iloc[:, 4] = df.iloc[:, 5]
            df.iloc[:, 5] = ""
            df = df.iloc[:, :-1]

            self.progress.emit(70)
            self.log.emit("📝 Đang tạo Word...")

            document = Document(template_path)
            table = document.tables[0]

            for i, row in enumerate(df.values, start=1):
                if i >= len(table.rows):
                    table.add_row()
                for j, value in enumerate(row):
                    cell = table.rows[i].cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            for i in range(len(table.rows) - 1, 0, -1):
                if not any(cell.text.strip() for cell in table.rows[i].cells):
                    table._element.remove(table.rows[i]._element)

            # CLOSE Word if output file is open
            close_word_if_file_open(self.output_path)

            document.save(self.output_path)
            time.sleep(1)
            subprocess.Popen(['start', self.output_path], shell=True)
            self.progress.emit(100)
            self.finished.emit(df)

        except Exception as e:
            self.log.emit(f"❌ Lỗi: {str(e)}")
            self.progress.emit(0)
            self.finished.emit(None)
        finally:
            browser.quit()

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(620, 720)
        MainWindow.setStyleSheet("""
            QWidget {
                background-color: #f1fafe;
                font-family: 'Segoe UI';
            }
            QPushButton {
                background-color: #007acc;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-size: 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #005b99;
            }
            QLineEdit, QTextEdit {
                background-color: #ffffff;
                border: 2px solid #d0eafc;
                border-radius: 8px;
                padding: 6px;
                font-size: 13px;
            }
            QLabel {
                color: #003366;
                font-size: 13pt;
            }
            QProgressBar {
                height: 24px;
                border-radius: 8px;
                background: #d0eafc;
            }
            QProgressBar::chunk {
                background-color: #007acc;
                border-radius: 8px;
                transition: all 0.5s ease-in-out;
            }
            QTabWidget::pane {
                border: 2px solid #d0eafc;
                border-radius: 8px;
                background: #ffffff;
                margin-top: 10px;
            }
            QTabBar::tab {
                background: #d0eafc;
                border: 2px solid #d0eafc;
                border-radius: 8px;
                padding: 8px 20px;
                font-size: 13px;
            }
            QTabBar::tab:selected {
                background: #007acc;
                color: white;
                font-weight: bold;
            }
            QTabBar::tab:hover {
                background: #aad4f5;
            }
        """)

        self.centralwidget = QtWidgets.QWidget(MainWindow)
        main_layout = QtWidgets.QVBoxLayout(self.centralwidget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)

        # New: Add a QTabWidget
        self.tabWidget = QtWidgets.QTabWidget()
        main_layout.addWidget(self.tabWidget)

        # First tab
        self.tab1 = QtWidgets.QWidget()
        tab1_layout = QtWidgets.QVBoxLayout(self.tab1)
        tab1_layout.setSpacing(15)

        label_mnv_layout = QtWidgets.QHBoxLayout()
        self.label_mnv = QtWidgets.QLabel("Nhập mã nhân viên:")
        self.mnv = QtWidgets.QLineEdit()
        self.mnv.setFont(QtGui.QFont("Segoe UI", 11))
        self.mnv.textChanged.connect(self.validate_input)
        label_mnv_layout.addWidget(self.label_mnv)
        label_mnv_layout.addWidget(self.mnv)

        btn_layout = QtWidgets.QHBoxLayout()
        self.create_btn = QtWidgets.QPushButton("Tạo phiếu xuất hàng")
        self.create_btn.clicked.connect(self.generate_report)
        self.clear_btn = QtWidgets.QPushButton("Clear")
        self.clear_btn.clicked.connect(self.clear_input)
        btn_layout.addWidget(self.create_btn)
        btn_layout.addWidget(self.clear_btn)

        self.progress_bar = QtWidgets.QProgressBar()
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setValue(0)

        self.log_output = QtWidgets.QTextEdit()
        self.log_output.setReadOnly(True)

        self.spinner_label = QtWidgets.QLabel()
        self.spinner_movie = QtGui.QMovie(resource_path("spinner.gif"))
        self.spinner_label.setMovie(self.spinner_movie)
        self.spinner_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.spinner_label.hide()

        # Add widgets to tab layout
        tab1_layout.addLayout(label_mnv_layout)
        tab1_layout.addLayout(btn_layout)
        tab1_layout.addWidget(self.progress_bar)
        tab1_layout.addWidget(self.log_output)
        tab1_layout.addWidget(self.spinner_label)

        self.tabWidget.addTab(self.tab1, "Tạo Phiếu Xuất Hàng")  

        self.tab2 = QtWidgets.QWidget()
        self.tabWidget.addTab(self.tab2, "In CTKM")

        MainWindow.setCentralWidget(self.centralwidget)


    def validate_input(self):
        text = self.mnv.text()
        if not text.isnumeric():
            self.mnv.setStyleSheet("border: 2px solid red;")
        else:
            self.mnv.setStyleSheet("border: 2px solid #d0eafc;")

    def on_report_finished(self, df):
        self.spinner_movie.stop()
        self.spinner_label.hide()
        if df is not None:
            self.log_output.append("🎉 Hoàn thành!")

    def generate_report(self):
        mnv = self.mnv.text()
        if not mnv:
            QtWidgets.QMessageBox.warning(None, "Lỗi", "Vui lòng nhập mã nhân viên")
            return

        self.log_output.clear()
        self.progress_bar.setValue(0)
        self.spinner_label.show()
        self.spinner_movie.start()

        output_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            None, "Lưu phiếu xuất hàng", "phieu_xuat.docx", "Word Documents (*.docx)")
        if not output_path:
            self.log_output.append("❌ Đã hủy lưu file.")
            self.spinner_movie.stop()
            self.spinner_label.hide()
            return

        self.thread = Worker(mnv, output_path)
        self.thread.progress.connect(self.progress_bar.setValue)
        self.thread.log.connect(self.log_output.append)
        self.thread.finished.connect(self.on_report_finished)
        self.thread.start()

    def clear_input(self):
        self.mnv.clear()
        self.progress_bar.setValue(0)
        self.log_output.clear()

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    latest = check_for_updates()
    if latest:
        prompt_update(latest)
    MainWindow = QtWidgets.QMainWindow()
    MainWindow.setWindowTitle("Phiếu Xuất Hàng - SCFP Tool")
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec())
